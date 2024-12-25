from docx import Document
from docx.shared import RGBColor
from nltk.sentiment.vader import SentimentIntensityAnalyzer
from nltk.tokenize import sent_tokenize
from datetime import date

import nltk

# Download lexicon
nltk.download('punkt')
nltk.download('vader_lexicon')

# Function to prepare natural language processing library and document writer
def result_preparation():
    # Create a new Document to write out a Sentiment Analysis report
    doc = Document()

    # Add a Title
    doc.add_heading('Sentiment Analysis Report', 0)

    # Add a sub-header
    header = doc.add_heading('', level=1)  # Empty heading text first

    # Add a run to the heading and set the text
    run = header.add_run("Overall Sentiment:")

    # Change the font color to black
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Create a SentimentIntensityAnalyzer object
    sia = SentimentIntensityAnalyzer()

    return sia, doc, run


# Function to analyze the sentiments of inputted user text
def sentiment_analyzer(input_text):

    # Do pre-preparation for writing the files
    sia, doc, run = result_preparation()

    # Step 1: Tokenize the text into sentences
    sentences = sent_tokenize(input_text)

    # Step 2: Initialize counters for slightly positive, overly positive, neutral, slightly negative and overly negative sentences
    slightly_positive_count = 0
    overly_positive_count = 0
    neutral_count = 0
    slightly_negative_count = 0
    overly_negative_count = 0

    # Step 3: Analyze each sentence's sentiment
    for sentence in sentences:
        sentiment_scores = sia.polarity_scores(sentence)
        compound_score = sentiment_scores['compound']

        # Classify each sentence's sentiment
        if compound_score >= 0.50:
            overly_positive_count += 1
        elif (compound_score >= 0.05) and (compound_score < 0.5):
            slightly_positive_count += 1
        elif (compound_score <= -0.05) and (compound_score > -0.5):
            slightly_negative_count += 1
        elif compound_score <= -0.50:
            overly_negative_count += 1
        else:
            neutral_count += 1

    # Step 5: Calculate the overall compound sentiment for confidence
    overall_sentiment = sia.polarity_scores(input_text)
    overall_confidence = abs(overall_sentiment['compound'])

    # Add a single paragraph
    para = doc.add_paragraph()

    # Get maximum polarity value and it's type
    max_key, max_value = max(list(overall_sentiment.items())[:3], key=lambda item: item[1])

    # Write polarity score in document
    run = para.add_run("Original Text: ")
    run.bold = True
    run = para.add_run(input_text).add_break()

    # Assign keywords as per polarity type
    if max_key == 'neg':
        max_key_new = 'Major Polarity Score (Negative): '
    elif max_key == 'neu':
        max_key_new = 'Major Polarity Score (Neutral): '
    else:
        max_key_new = 'Major Polarity Score (Positive): '

    # Write Overall Polarity Score
    run = para.add_run(max_key_new)
    run.bold = True
    run = para.add_run(str(abs((max_value)*100)) + '%').add_break()

    # Write Confidence Value
    run = para.add_run("Overall Confidence: ")
    run.bold = True
    run = para.add_run(str(overall_confidence)).add_break()

    # Write Other Polarity scores
    run = para.add_run("Other minor polarity scores of the text are ")

    # Iterate through the first three key-value pairs of polarity scores
    for key, value in list(overall_sentiment.items())[:3]:

        # Write other polarity values
        if str({key})[2:-2] != max_key:

            if str({key})[2:-2] == 'neg':
                run = para.add_run(" negative with ")
                run = para.add_run(str({round(value*100, 2)})[1:-1] +'%')
                run.bold = True
            elif str({key})[2:-2] == 'neu':
                run = para.add_run(" and neutral with ")
                run = para.add_run(str({round(value*100, 2)})[1:-1] +'%')
                run.bold = True
            elif str({key})[2:-2] == 'pos':
                run = para.add_run(" and positive with ")
                run = para.add_run(str({round(value*100, 2)})[1:-1] +'%')
                run.bold = True
        else:
            pass

    # Put a horizontal break
    run.add_break()

    # Write statement declaring individual polarities of individual statements
    run = para.add_run("Individual polarity of the sentences are as follows: ").add_break()

    # Counting Variable
    counting_variable = 1

    # Iterate through each sentence in the user inputted text
    for sentence in sentences:

        # Get polarity scores and confidence interval of the sentiment of each sentence
        sentiment_scores = sia.polarity_scores(sentence)
        compound_score = sentiment_scores['compound']

        # Get the maximum polarity value and the type
        max_key, max_value = max(list(sentiment_scores.items())[:3], key=lambda item: item[1])

        # Write the sentences in the word file
        para = doc.add_paragraph()
        run = para.add_run("Sentence " + str(counting_variable) + ": ")
        run.bold = True
        run = para.add_run(sentence).add_break()

        # Assign explanatory words for different polarities
        if max_key == 'neg':
            max_key_new = 'Overall Polarity Score (Negative): '
        elif max_key == 'neu':
            max_key_new = 'Overall Polarity Score (Neutral): '
        elif max_key == 'pos':
            max_key_new = 'Overall Polarity Score (Positive): '

        # Write polarity scores for individual sentences
        run = para.add_run(max_key_new)
        run.bold = True
        run = para.add_run(str(round(max_value * 100, 2)) + '%').add_break()

        # Write confidence scores for individual sentence's polarities
        run = para.add_run("Overall Confidence: ")
        run.bold = True
        run = para.add_run(str(round(abs(compound_score), 2))).add_break()

        # Increment counting variable
        counting_variable += 1

    # Save the document to a file
    doc.save('Result Files/Sentiment Analysis Report.docx')

# sentiment_analyzer("Everything feels pointless. No matter how hard I try, it all seems like a futile effort, a "
#                    "never-ending loop of disappointment. The world is full of noise, yet there’s no substance, no "
#                    "real meaning behind any of it. Every day is just a repeat of the last — the same dull routine, "
#                    "the same empty conversations, the same relentless, crushing weight of existence. Nothing improves, "
#                    "nothing changes. Hope is just a fleeting illusion, something to keep us moving forward even when "
#                    "we know deep down it’s all an empty promise. People talk, but they never listen. They pretend to "
#                    "care, but their words are hollow, meaningless. Every attempt to connect feels like a failure, a "
#                    "reminder that we’re all just drifting in this vast, indifferent universe, unable to truly reach "
#                    "one another. The more we try to understand, the less we really know.And all the things I once "
#                    "cared about? They've faded into the background, forgotten in the chaos. What’s the point of "
#                    "ambition or dreams when they’re just as fragile as everything else? Life just keeps going, "
#                    "whether you want it to or not, and every step forward feels like a small surrender, a little piece "
#                    "of your soul chipped away by the weight of the world. It’s hard to escape the feeling that nothing "
#                    "is ever truly right. There’s always a dark cloud looming, no matter how hard you try to push it "
#                    "away. The good moments are fleeting, and the bad ones seem to stretch on forever. It’s exhausting. "
#                    "And in the end, does any of it even matter?")