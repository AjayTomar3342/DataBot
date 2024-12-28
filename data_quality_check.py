import pandas as pd
from docx import Document
from docx.shared import RGBColor

# Create a new Document to write out a Data Quality Check report
doc = Document()

# Add a Title
doc.add_heading('Data Quality Report', 0)

# Variable to store type of information
information_type = ""

# Variable to store solution
solution = ""

# Function to write the comments (on dataframe level) into final Data Quality Report
def data_quality_report(heading, comment,  type, solution):

    # Add a sub-header
    header = doc.add_heading('', level=1)  # Empty heading text first

    # Add a run to the heading and set the text
    run = header.add_run(heading)

    # Change the font color to black
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Add a single paragraph
    para = doc.add_paragraph()

    # Case when statement is of type Warning
    if type == 'warning':
        # Write Warning Statement
        run = para.add_run("WARNING: ")
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color using RGB
        run = para.add_run(comment)

        # Write Solution Statement
        run = para.add_run('\n')  # This simulates a line break
        run = para.add_run("SOLUTION: ")
        run.font.color.rgb = RGBColor(144, 238, 144)  # Red color using RGB
        run = para.add_run(solution)


    # Case when statement is of type Information
    if type == 'information':
        run = para.add_run("INFO: ")
        run = para.add_run(comment)

    # Save the document to a file
    doc.save('Result Files/Data Quality Report.docx')

# Function to write the comments (on columns level) into final Data Quality Report
def data_quality_report_columns(heading, comment,  type, solution, column, first_time):

    if type == "remarks":
        # Add a single paragraph
        para = doc.add_paragraph(' ')

        run = para.add_run(solution[:8])
        run.font.color.rgb = RGBColor(144, 238, 144)  # Red color using RGB
        run = para.add_run(solution[8:])


    if column == "yes":

        # Add a single paragraph
        para = doc.add_paragraph(' ')

        # Case when statement is of type Warning
        if type == 'warning':
            if column == "yes":
                log_code = "• "
                color_code = RGBColor(0, 0, 0)
            else:
                log_code = "WARNING: "
                color_code = RGBColor(255, 0, 0)


            # Write Warning Statement
            run = para.add_run(log_code)
            run.font.color.rgb = color_code  # Red color using RGB
            run = para.add_run(comment)

            # Write Solution Statement
            if solution != "":
                run = para.add_run('\n')  # This simulates a line break
                run = para.add_run('\n')  # This simulates a line break
            else:
                pass
            run = para.add_run(solution[:9])
            run.font.color.rgb = RGBColor(144, 238, 144)  # Red color using RGB
            run = para.add_run(solution[9:])

        # Case when statement is of type Information
        if type == 'information':
            if column == "yes":
                log_code = "• "
            else:
                log_code = "INFO: "
            run = para.add_run(log_code)
            run = para.add_run(comment)

    elif column == "no":
        # Add a sub-header
        header = doc.add_heading('', level=1)  # Empty heading text first

        # Add a run to the heading and set the text
        run = header.add_run(heading)

        # Change the font color to black
        run.font.color.rgb = RGBColor(0, 0, 0)

        if type == "information":
            # Add a single paragraph
            para = doc.add_paragraph(' ')
            run = para.add_run(comment)
        else:
            # Add a single paragraph
            para = doc.add_paragraph(' ')
            run = para.add_run(comment[:9])
            run.font.color.rgb = RGBColor(255, 0, 0)
            run = para.add_run(comment[9:])

    # Save the document to a file
    doc.save('Result Files/Data Quality Report.docx')

# Function to get details about unnamed columns
def unnamed_columns_count(df):

    solution = "Assign column names manually so as to remove confusion during data file handling."

    # Variable to store count of unnamed columns
    count_of_unnamed_columns = len(df.columns[df.columns.str.contains("Unnamed")])

    # If there are no unnamed columns, then information type is just information
    if count_of_unnamed_columns == 0:
        information_type = "information"
    # If the above case is not true, then information type is warning
    else:
        information_type = "warning"
        solution = "Assign column names manually so as to remove confusion during data file handling."

    # Write a comment about Unnamed Dataframe columns in Data Quality Report
    data_quality_report( "Unnamed Columns Count (Dataset Level): ", "There are " + str(count_of_unnamed_columns)
                         + " columns found which had no name.", information_type, solution)

    # Save the document to a file
    doc.save('Result Files/Data Quality Report.docx')


# Function to get duplicate rows percentage
def duplicate_rows_percentage(df):

    # Identify duplicate rows (excluding the first occurrence)
    duplicates = df.duplicated()

    # Calculate the number of duplicate rows (True means duplicate)
    num_duplicates = duplicates.sum()

    # Calculate total rows
    total_rows = len(df)

    # Calculate the percentage of duplicate rows
    duplicate_percentage = (num_duplicates / total_rows) * 100

    # If there are no duplicate rows, then information type is just information
    if duplicate_percentage == 0:
        information_type = "information"
        solution = ""
    # If the above case is not true, then information type is warning
    else:
        information_type = "warning"
        solution = ("SOLUTION: Remove duplicates using a duplication removal tools, or you can use pandas "
                    "function (drop_duplicates) to remove them.")

    # Write a comment about Duplicate Rows percentage in Data Quality Report
    data_quality_report( "Duplicate Rows Percentage (Dataset Level): ", "There are " + str(duplicate_percentage)
                         + "% duplicate rows in the uploaded dataset.", information_type, solution)

    # Save the document to a file
    doc.save('Result Files/Data Quality Report.docx')


# Function to identify outliers in a DataFrame using IQR
def detect_outliers(df):
    outliers = {}
    for column in df.select_dtypes(include=['float64', 'int64']).columns:
        # Calculate the IQR for each column
        Q1 = df[column].quantile(0.25)
        Q3 = df[column].quantile(0.75)
        IQR = Q3 - Q1

        # Determine the lower and upper bounds
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR

        # Identify the outliers
        outlier_values = df[(df[column] < lower_bound) | (df[column] > upper_bound)]

        if not outlier_values.empty:
            outliers[column] = outlier_values

    # If there are columns containing outliers, then information type is just information
    if len(outliers) == 0:
        information_type = "information"
    # If the above case is not true, then information type is warning
    else:
        information_type = "warning"
        solution = ""

    # Write a comment about Outliers inside columns in Data Quality Report
    data_quality_report_columns( "Outliers Detection (Column Level): ", "WARNING: There are " + str(len(outliers))
                         + " columns contain outlier values. They are as follows: ", information_type, solution,
                                 "no", "true")

    # If outliers columns are not equal to 0
    if len(outliers) != 0:

        counting_variable = 0

        # Iterate through columns containing outliers
        for column, outlier_data in outliers.items():

            # Assign information to the type variable
            information_type = "information"

            # If last column has been reached, give the solution too
            if counting_variable == len(outliers) - 1:
                solution = ("SOLUTION: Outliers in some columns could be an issue as they might impact statistical measures such "
                            "as mean and standard deviation in a wrong way so that they show influenced results. Moreover, if "
                            "one plans to use the outlier affected columns in regression models, then outliers can distort the "
                            "relationship between variables in the regression model and impact the model's coefficients and"
                            " overall fit disproportionately. If the values in the range above are same, then it might be due to a singular "
                            "value being the sole outlier but having high frequency in the column.")

                # Assign information to the type warning
                information_type = "warning"

            # Write a comment about Outliers inside columns in Data Quality Report
            data_quality_report_columns( "", "Column " + str({column})[2:-2] + " contains " + str(outlier_data.shape[0])
                                    + " outlier values ranging from " + str(int(outlier_data[column].min())) + " to "
                                    + str(int(outlier_data[column].max())) + ".", information_type, solution, "yes",
                                         "false")

            # Increment counting variable
            counting_variable+=1

    # Save the document to a file
    doc.save('Result Files/Data Quality Report.docx')

# Function to null values percentage in each column
def null_values_proportion(df):

    # Dictionary to keep column names and their na values proportion
    proportions = {}

    # Iterate through the dataframe
    for column in df.columns:
        # Save column names in the above-mentioned dictionary
        proportions.setdefault('columns', []).append(column)
        null_count = df[column].isnull().sum()  # Count of nulls
        total_count = len(df)  # Total number of rows
        # Save na values proportions for each column in the above-mentioned dictionary
        proportions.setdefault('percentages', []).append(round((null_count/total_count)*100,2))

    # Create lists to store column names and their respective na values proportions
    column_names = list(df.columns)
    na_proportions = list(proportions['percentages'])


    # Check if there are null values in any column of the dataframe
    if sum(proportions['percentages']) > 0:
        solution = ""
        # Write a comment about Null values of dataframe columns in Data Quality Report
        data_quality_report_columns("Null Values Proportion (Column Level): ", "WARNING: The Null Values "
                                    "Proportion of all columns are as follows: ", "warning", solution,
                                    "no", "true")
        # Variable to count flow of a for loop
        counting_value = 0

        # Iterate through column names and their respective na values proportions
        for i in range(0, len(column_names)):

            # Check if data is being written about last column of the dataframe
            if counting_value == len(column_names) - 1:

                # Assign solution value
                solution = ("SOLUTION: A high count of Na values inside a column might lead to ineffective analysis "
                            "results and therefore na's have to be replaced either by a mean of the non-na values of "
                            "the column or a manual value which makes business or practical sense")
            else:
                pass
            # Write a comment about Null values proportion of dataframe columns in Data Quality Report
            data_quality_report_columns("Null Values Proportion (Column Level): ", "Column "
                                        + str(column_names[i]) + " contains " + str(na_proportions[i]) +
                                        "% of Na Values","warning", solution, "yes", "true")

            # Increment counting value variable
            counting_value += 1

    else:
        solution = ""
        # Write a comment about Null values of dataframe columns in Data Quality Report
        data_quality_report_columns("Null Values Proportion (Column Level): ", "INFO: There are no null "
                                    "values in the uploaded dataset.", "information", solution,
                                    "no", "true")

    # Save the document to a file
    doc.save('Result Files/Data Quality Report.docx')


# Function to get data type consistency of each column
def check_columns_data_type_consistency(df):


    # Dictionary to store columns which have inconsistent data types
    inconsistent_columns = {}

    # Iterate through the data type
    for column in df.columns:

        # Get the unique types found in the column
        column_types = df[column].apply(type).unique()

        # If more than one unique type is found, it indicates inconsistency
        if len(column_types) > 1:
            inconsistent_columns[column] = column_types


    # Write about Inconsistent data types inside all relevant columns in the Data Quality Report
    if inconsistent_columns:

        solution = ""
        # Write a comment about Inconsistent data types of dataframe columns in Data Quality Report
        data_quality_report_columns("Inconsistent  Data Types Check (Column Level): ", "WARNING: There are "
                                    + str(len(inconsistent_columns.keys())) + (" columns with inconsistent data type "
                                    "values. They are as follows: "), "warning", solution, "no", "true")

        # Variable to count flow of a for loop
        counting_value = 0

        # Iterate through columns having inconsistent data types in their values
        for column, types in inconsistent_columns.items():

            # Check if data is being written about last column of the dataframe
            if counting_value == len(inconsistent_columns) - 1:

            # Assign solution value
                solution = ("SOLUTION: A single column cannot contain values of multiple data types. Use pandas functions"
                            " such as astype() or pd.to_datetime() or pd.to_numeric() to coerce each relevant column "
                            "values into a single data type. The next section of this report calculates summary statistics"
                            " and if one of the data inconsistency affected column is supposed to be integer or float "
                            "type, then it won't show up in the summary statistics report. This is one such example of"
                            " the disadvantages of a column having values of multiple data types.")

                # Write a comment about Inconsistent dataframe columns in Data Quality Report
                data_quality_report_columns("Inconsistent  Data Types Check (Column Level): ", "Column "
                                            + str(column) + " contains multiple data type values, namely "
                                            + str(types[0])[8:-2] + " and " + str(types[1])[8:-2] +
                                            ".", "warning", solution, "yes", "true")

            else:


                # Write a comment about Inconsistent dataframe columns in Data Quality Report
                data_quality_report_columns("Inconsistent  Data Types Check (Column Level): ", "Column "
                                            + str(column) + " contains multiple data type values, namely "
                                                + str(types[0])[8:-2] + " and " + str(types[1])[8:-2] +
                                            ".", "warning", solution, "yes", "true")

            counting_value += 1

    else:
        solution = ""
        # Write a comment about Inconsistent data types of columns in Data Quality Report
        data_quality_report_columns("Inconsistent  Data Types Check (Column Level): ", "INFO: There are no "
                                    "columns which have inconsistent data types.", "information", solution,
                                    "no", "true")

    # Save the document to a file
    doc.save('Result Files/Data Quality Report.docx')


# Function to get summary statistics of each column
def summary_statistics(df):

    # Get all integer and float columns
    df = df.select_dtypes(include=['int64', 'float64'])


    # Check if the dataframe is empty
    if df.empty == False:
        solution = ""
        # Write a comment about Summary statistics dataframe columns in Data Quality Report
        data_quality_report_columns("Summary Statistics (Column Level): ", "INFO: Summary statistics "
                                    "for the number-related columns are as follows:",
                                    "information", solution, "no", "true")

        # Variable to count flow of a for loop
        counting_value = 0

        # Iterate through the data frame columns
        for column in df.columns:
            # Check if data is being written about last column of the dataframe
            if counting_value == df.shape[1] - 1:
                # Assign solution value
                solution = (
                    "SOLUTION: If you find insufficient number-related columns above then please check if"
                    " some columns which should be number-related have multiple data types values inside them (as "
                    "mentioned the above section). Please make sure no such column is being considered as a multiple"
                    " data type column.")

                try:
                    # Write a comment about Summary statistics of columns in Data Quality Report
                    data_quality_report_columns("Summary Statistics (Column Level): ", "Column " + str(column) +
                                                " has a non-null row count of " + str(int(df[column].count())) +
                                                ", mean of " + str(int(df[column].mean())) + ", median " +
                                                 str(int(df[column].median())) + ", mode " + str(int(df[column].mode())) +
                                                ", minimum value " + str(int(df[column].min())) + ", maximum value "
                                                + str(int(df[column].max())) + ", standard deviation "
                                                + str(int(df[column].std())) + ", variance " + str(int(df[column].var())) +
                                                ", range " + str(int(df[column].max()) - int(df[column].min())) +
                                                ", first quartile " + str(int(df[column].quantile(0.25))) +
                                                ", third quartile " + str(int(df[column].quantile(0.75))) +
                                                ", skewness " + str(int(df[column].skew())) +
                                                ", kurtosis " + str(int(df[column].kurt())) +
                                                ", and has " + str(int(df[column].isna().sum())) + " na values."
                                                , "warning", solution, "yes", "true")
                except Exception as e:
                    # Code to handle the exception
                    df.drop(column, axis=1, inplace=True)

            else:
                try:
                    # Write a comment about Summary statistics of columns in Data Quality Report
                    data_quality_report_columns("Summary Statistics (Column Level): ", "Column " + str(column) +
                                                " has a non-null row count of " + str(int(df[column].count())) +
                                                ", mean of " + str(int(df[column].mean())) + ", median " +
                                                 str(int(df[column].median())) + ", mode " + str(int(df[column].mode())) +
                                                ", minimum value " + str(int(df[column].min())) + ", maximum value "
                                                + str(int(df[column].max())) + ", standard deviation "
                                                + str(int(df[column].std())) + ", variance " + str(int(df[column].var())) +
                                                ", range " + str(int(df[column].max()) - int(df[column].min())) +
                                                ", first quartile " + str(int(df[column].quantile(0.25))) +
                                                ", third quartile " + str(int(df[column].quantile(0.75))) +
                                                ", skewness " + str(int(df[column].skew())) +
                                                ", kurtosis " + str(int(df[column].kurt())) +
                                                ", and has " + str(int(df[column].isna().sum())) + " na values."
                                                , "warning", solution, "yes", "true")

                except Exception as e:
                    # Code to handle the exception
                    df.drop(column, axis=1, inplace=True)



            # Increment the counting variable
            counting_value += 1



    else:

        solution = ""
        # Write a comment about insufficient number data types columns in Data Quality Report
        data_quality_report_columns("Summary Statistics (Column Level): ", "INFO: There are no "
                                    "number-related columns in the dataset and therefore summary statistics cannot "
                                    "be calculated. It may be due to the lack of relevant number-related columns or "
                                    "number-related columns having multiple data types in their rows.",
                                    "information", solution,"no", "true")

    # Save the document to a file
    doc.save('Result Files/Data Quality Report.docx')


# Function to add final comments to the file
def final_comments(df):
    # Assign remark value
    remarks = ("REMARKS: If you find insufficient number-related columns above then please check if"
                " some columns which should be number-related have multiple data types values inside them (as "
                "mentioned the above section). Please make sure no such column is being considered as a multiple"
                " data type column.")
    # Write a comment about Summary statistics of columns in Data Quality Report
    data_quality_report_columns("Summary Statistics (Column Level): ", "",
                                "remarks", remarks, "yes", "true")

    # Save the document to a file
    doc.save('Result Files/Data Quality Report.docx')

# Function to do data quality analysis
def data_quality_check():

    # Read the user uploaded file
    df = pd.read_csv('User Files/data_quality_check_df.csv')

    # Delete first column and three rows
    df = df.drop(df.columns[0], axis=1)
    df = df.iloc[3:]

    # Use a function to get count of unnamed columns
    unnamed_columns_count(df)

    # Use a function to get percentages of duplicate rows
    duplicate_rows_percentage(df)

    # Use a function to get outliers of dataframe
    detect_outliers(df)

    # Use a function to get null-values of dataframe
    null_values_proportion(df)

    # Use a function to find data type consistency of each column of the dataframe
    check_columns_data_type_consistency(df)

    # Use a function to find summary statistics of each column of the dataframe
    summary_statistics(df)

    # Use a function to give last comments to the file
    final_comments(df)

    # Read the resulting docx file
    doc = Document('Result Files/Data Quality Report.docx')

    # print(doc)

    return doc

    # print(df)

# data_quality_check()



